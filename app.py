import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from datetime import datetime
from ydata_profiling import ProfileReport
import nbformat
from nbformat.v4 import new_notebook, new_code_cell
from io import BytesIO
import os
from docx import Document
from docx.shared import Inches
from xhtml2pdf import pisa

# ===== PAGE CONFIGURATION =====
st.set_page_config(
    page_title="Data Analysis Dashboard",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for styling
st.markdown("""
<style>
    .main {
        background-color: #f8f9fa;
    }
    .st-emotion-cache-1v0mbdj {
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    .st-bq {
        border-left: 3px solid #4e89e5;
    }
    .metric-box {
        background-color: white;
        border-radius: 10px;
        padding: 15px;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        text-align: center;
    }
    .stPlotlyChart {
        border-radius: 10px;
    }
    .download-btn {
        margin-bottom: 10px;
        width: 100%;
    }
</style>
""", unsafe_allow_html=True)

# ===== DATA LOADING =====
@st.cache_data
def load_sample_data():
    """Load built-in sample dataset"""
    data = {
        'CustomerID': range(1, 101),
        'Age': np.random.randint(18, 70, size=100),
        'Gender': np.random.choice(['Male', 'Female'], size=100),
        'PurchaseAmount': np.random.normal(100, 30, size=100).round(2),
        'Category': np.random.choice(['Electronics', 'Clothing', 'Home', 'Food'], size=100),
        'Rating': np.random.randint(1, 6, size=100)
    }
    return pd.DataFrame(data)

@st.cache_data
def load_user_data(uploaded_file):
    """Load user-uploaded data with maximum compatibility"""
    try:
        # Reset file pointer to beginning
        uploaded_file.seek(0)
        
        # First try to detect if it's Excel
        if uploaded_file.name.lower().endswith(('.xlsx', '.xls')):
            try:
                return pd.read_excel(uploaded_file, engine='openpyxl')
            except:
                try:
                    return pd.read_excel(uploaded_file, engine='xlrd')
                except:
                    st.error("Excel file loading failed. Please ensure you have openpyxl and xlrd installed.")
                    return None
        
        # Handle CSV and other text files
        else:
            # Try common encodings in order of likelihood
            encodings = ['utf-8', 'latin1', 'windows-1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    return pd.read_csv(uploaded_file, encoding=encoding)
                except UnicodeDecodeError:
                    continue
            
            # Final fallback with error replacement
            try:
                uploaded_file.seek(0)
                return pd.read_csv(uploaded_file, encoding='utf-8', errors='replace')
            except Exception as e:
                st.error(f"Failed to load file: {str(e)}")
                return None
                
    except Exception as e:
        st.error(f"Unexpected error loading file: {str(e)}")
        return None

# Initialize session state
if 'df' not in st.session_state:
    st.session_state.original_df = load_sample_data()
    st.session_state.df = st.session_state.original_df.copy()
    st.session_state.data_source = "sample"
    st.session_state.filters_applied = False
    st.session_state.filter_params = {}  
    st.session_state.current_chart = None
    st.session_state.chart_type = None
    st.session_state.x_axis = None
    st.session_state.y_axis = None
    st.session_state.saved_charts = []

# ===== OPTIMIZED FILTERING =====
def apply_filters():
    """Apply filters only to columns that have active filter values"""
    if not st.session_state.get('filter_params'):
        st.session_state.df = st.session_state.original_df.copy()
        st.session_state.filters_applied = False
        return
    
    df = st.session_state.original_df.copy()
    filter_params = st.session_state.filter_params
    any_filters_applied = False
    
    for col, filter_value in filter_params.items():
        col_type = st.session_state.original_df[col].dtype
        
        # Handle numeric filters
        if isinstance(filter_value, tuple) and np.issubdtype(col_type, np.number):
            min_val, max_val = filter_value
            if min_val is not None and max_val is not None:
                df = df[(df[col] >= min_val) & (df[col] <= max_val)]
                any_filters_applied = True
        
        # Handle categorical filters
        elif isinstance(filter_value, list):
            if filter_value:  # Only apply if values are selected
                df = df[df[col].isin(filter_value)]
                any_filters_applied = True
        
        # Handle date filters
        elif pd.api.types.is_datetime64_any_dtype(col_type):
            start_date, end_date = filter_value
            if start_date and end_date:
                df = df[
                    (df[col] >= pd.to_datetime(start_date)) & 
                    (df[col] <= pd.to_datetime(end_date))
                ]
                any_filters_applied = True
    
    st.session_state.df = df
    st.session_state.filters_applied = any_filters_applied

# ===== SIDEBAR =====
with st.sidebar:  # <-- This creates the sidebar container
    # Logo and Data Input Section
    st.image("assets/logo.png", width=120)
    st.title("Data Input")
    
    # File uploader (must be inside the sidebar context)
    uploaded_file = st.file_uploader("Upload your data file", type=["csv", "xlsx"])
    if uploaded_file is not None:
        user_data = load_user_data(uploaded_file)
        if user_data is not None:
            # Only clear saved_charts if this is a genuinely new dataset
            if (st.session_state.data_source == "sample" or 
                not st.session_state.original_df.equals(user_data)):
                st.session_state.saved_charts = []  # Clear saved charts only for new data
            
            st.session_state.original_df = user_data
            st.session_state.df = st.session_state.original_df.copy()
            st.session_state.data_source = "user"
            st.session_state.filters_applied = False
            st.session_state.filter_params = {}
            st.success("Data loaded successfully!")
    else:
        if st.session_state.data_source == "sample":
            st.info("Using sample dataset")
    
    st.divider()
    
    # Data Status Section
    st.subheader("Data Status")
    if st.session_state.filters_applied:
        st.success("✅ Filtered data")
        if st.button("❌ Reset All Filters", key="reset_filters"):
            st.session_state.df = st.session_state.original_df.copy()
            st.session_state.filters_applied = False
            st.session_state.filter_params = {}
            st.rerun()
    else:
        st.info("🔄 Original data")
    
    st.divider()
    
    # Filter Controls Section
    st.header("Filter Controls")
    filter_enabled = st.checkbox("Enable Filters", value=True, key="enable_filters")
    
    if filter_enabled:
        if 'filter_params' not in st.session_state:
            st.session_state.filter_params = {}
        
        all_columns = st.session_state.original_df.columns.tolist()
        
        for col in all_columns:
            with st.expander(f"Filter: {col}", expanded=False):
                col_type = st.session_state.original_df[col].dtype
                
                if np.issubdtype(col_type, np.number):
                    min_val = float(st.session_state.original_df[col].min())
                    max_val = float(st.session_state.original_df[col].max())
                    
                    # Initialize with None instead of min/max
                    current_min, current_max = st.session_state.filter_params.get(col, (None, None))
                    
                    new_min = st.number_input(
                        "Minimum value",
                        min_value=min_val,
                        max_value=max_val,
                        value=current_min,
                        key=f"min_{col}",
                        placeholder="Enter min value"
                    )
                    new_max = st.number_input(
                        "Maximum value",
                        min_value=min_val,
                        max_value=max_val,
                        value=current_max,
                        key=f"max_{col}",
                        placeholder="Enter max value"
                    )
                    
                    # Only store if both values are provided
                    if new_min is not None and new_max is not None:
                        st.session_state.filter_params[col] = (new_min, new_max)
                    else:
                        st.session_state.filter_params.pop(col, None)
                
                elif col_type == 'object' or pd.api.types.is_categorical_dtype(col_type):
                    options = st.session_state.original_df[col].unique().tolist()
                    current_filter = st.session_state.filter_params.get(col, [])
                    
                    selected = st.multiselect(
                        "Select values to include",
                        options,
                        default=current_filter,
                        key=f"select_{col}",
                        placeholder="Select values..."
                    )
                    
                    # Only store if values are selected
                    if selected:
                        st.session_state.filter_params[col] = selected
                    else:
                        st.session_state.filter_params.pop(col, None)
                
                elif pd.api.types.is_datetime64_any_dtype(col_type):
                    min_date = st.session_state.original_df[col].min().to_pydatetime()
                    max_date = st.session_state.original_df[col].max().to_pydatetime()
                    current_filter = st.session_state.filter_params.get(col, (None, None))
                    
                    dates = st.date_input(
                        "Select date range",
                        value=current_filter,
                        min_value=min_date,
                        max_value=max_date,
                        key=f"date_{col}"
                    )
                    
                    # Only store if both dates are selected
                    if len(dates) == 2:
                        st.session_state.filter_params[col] = dates
                    else:
                        st.session_state.filter_params.pop(col, None)
        
        # Add explicit apply button
        apply_col, clear_col = st.columns(2)
        with apply_col:
            if st.button("Apply Filters", key="apply_filters"):
                apply_filters()
        with clear_col:
            if st.button("Clear Filters", key="clear_filters"):
                st.session_state.df = st.session_state.original_df.copy()
                st.session_state.filters_applied = False
                st.session_state.filter_params = {}
                st.rerun()
# ===== MAIN DASHBOARD =====
st.title("📊 Data Analysis Dashboard")
st.markdown("Interactive tool for exploring datasets and generating insights")

# Processing Summary
st.header("Processing Summary")
col1, col2 = st.columns(2)
with col1:
    original_rows = len(st.session_state.original_df)
    current_rows = len(st.session_state.df)
    st.metric("Rows Remaining", current_rows, delta=f"{- (original_rows - current_rows)} filtered")

with col2:
    original_cols = len(st.session_state.original_df.columns)
    current_cols = len(st.session_state.df.columns)
    if current_cols < original_cols:
        st.metric("Columns Remaining", current_cols, delta="Columns removed", delta_color="inverse")
    else:
        st.metric("Columns", current_cols)

# Key Metrics
st.header("Key Metrics")
metric_cols = st.columns(4)
metric_data = [
    ("Total Records", len(st.session_state.df)),
    ("Columns", len(st.session_state.df.columns)),
    ("Missing Values", st.session_state.df.isnull().sum().sum()),
    ("Duplicates", st.session_state.df.duplicated().sum())
]

for i, (title, value) in enumerate(metric_data):
    with metric_cols[i]:
        st.markdown(f"""
        <div class="metric-box">
            <h3 style="color:#4e89e5">{value}</h3>
            <p>{title}</p>
        </div>
        """, unsafe_allow_html=True)

# Data Preview
st.divider()
st.header("Data Preview")
if st.checkbox("Show raw data", True):
    st.dataframe(
        st.session_state.df,
        height=400,
        use_container_width=True,
        hide_index=True
    )

# Automated Insights
st.divider()
st.header("Automated Insights")
if st.button("Generate Insights", key="insights_btn"):
    with st.spinner("Analyzing data..."):
        insights = []
        df = st.session_state.df
        
        insights.append(f"**Dataset contains {len(df)} records** with {len(df.columns)} features")
        dtype_counts = df.dtypes.value_counts()
        insights.append(f"**Data types:** {', '.join([f'{count} {dtype}' for dtype, count in dtype_counts.items()])}")
        
        numeric_cols = df.select_dtypes(include=np.number).columns
        if len(numeric_cols) > 0:
            insights.append("**Numeric Features:**")
            for col in numeric_cols:
                stats = {
                    "min": df[col].min(),
                    "max": df[col].max(),
                    "mean": df[col].mean(),
                    "median": df[col].median()
                }
                insights.append(
                    f"- {col}: Min={stats['min']:.2f}, Max={stats['max']:.2f}, "
                    f"Mean={stats['mean']:.2f}, Median={stats['median']:.2f}"
                )
        
        cat_cols = df.select_dtypes(include=['object', 'category']).columns
        if len(cat_cols) > 0:
            insights.append("**Categorical Features:**")
            for col in cat_cols:
                top_values = df[col].value_counts().nlargest(3)
                insights.append(f"- {col}: {', '.join([f'{val} ({count})' for val, count in top_values.items()])}")
        
        if len(numeric_cols) >= 2:
            corr = df[numeric_cols].corr().unstack().sort_values(ascending=False)
            high_corr = corr[(abs(corr) > 0.5) & (corr < 1)].drop_duplicates().head(2)
            if not high_corr.empty:
                insights.append("**Strong Correlations:**")
                for (col1, col2), value in high_corr.items():
                    insights.append(f"- {col1} & {col2}: {value:.2f} correlation")
        
        with st.container():
            for insight in insights:
                st.markdown(insight)
            st.success("Analysis complete!")

# Interactive Visualizations
st.divider()
st.header("Interactive Visualizations")
if st.session_state.filters_applied:
    st.warning(f"""
    🔍 Visualizations show PROCESSED data. 
    - Removed {len(st.session_state.original_df) - len(st.session_state.df)} rows 
    ({(len(st.session_state.original_df) - len(st.session_state.df)) / len(st.session_state.original_df):.1%} reduction)
    """)

col1, col2 = st.columns(2)
with col1:
    chart_type = st.selectbox(
        "Chart Type",
        ["Bar Chart", "Histogram", "Scatter Plot", "Pie Chart", "Box Plot"],
        index=0
    )

with col2:
    x_axis = st.selectbox(
        "X-axis",
        st.session_state.df.columns,
        index=0
    )
    
    y_axis = st.selectbox(
        "Y-axis (if applicable)",
        ["None"] + st.session_state.df.columns.tolist(),
        index=0
    )
    if y_axis == "None":
        y_axis = None

st.subheader(f"{chart_type} Visualization")
try:
    fig = None
    df = st.session_state.df
    
    if chart_type == "Bar Chart":
        if y_axis and y_axis != "None":
            fig = px.bar(df, x=x_axis, y=y_axis, title=f"{y_axis} by {x_axis}", color=x_axis)
        else:
            count_df = df[x_axis].value_counts().reset_index()
            count_df.columns = [x_axis, 'count']
            fig = px.bar(count_df, x=x_axis, y='count', title=f"Count of {x_axis}", color=x_axis)
    
    elif chart_type == "Histogram":
        bins = st.slider("Number of bins", 5, 100, 20, key='hist_bins')
        if df[x_axis].dtype in ['int64', 'float64']:
            fig = px.histogram(
                df, 
                x=x_axis, 
                nbins=bins,
                title=f"Distribution of {x_axis}",
                marginal='box',
                color_discrete_sequence=px.colors.qualitative.Plotly
            )
        else:
            st.warning("Histogram requires numeric data. Please select a numeric column.")
    
    elif chart_type == "Pie Chart":
        count_df = df[x_axis].value_counts().reset_index()
        count_df.columns = ['category', 'count']
        fig = px.pie(
            count_df, 
            names='category', 
            values='count', 
            title=f"Distribution of {x_axis}",
            hole=0.3,
            color_discrete_sequence=px.colors.qualitative.Plotly
        )
    
    elif chart_type == "Scatter Plot":
        if y_axis and y_axis != "None":
            fig = px.scatter(
                df, 
                x=x_axis, 
                y=y_axis, 
                title=f"{y_axis} vs {x_axis}",
                trendline='ols',
                color=x_axis
            )
        else:
            st.warning("Please select both X and Y axes for scatter plot")
    
    elif chart_type == "Box Plot":
        if y_axis and y_axis != "None":
            fig = px.box(
                df, 
                x=x_axis, 
                y=y_axis, 
                title=f"{y_axis} Distribution by {x_axis}",
                color=x_axis,
                color_discrete_sequence=px.colors.qualitative.Plotly
            )
        else:
            st.warning("Please select both X and Y axes for box plot")
    
    if fig:
        fig.update_layout(
            plot_bgcolor='rgba(0,0,0,0)',
            paper_bgcolor='rgba(0,0,0,0)',
            font=dict(color="#2c3e50"),
            colorway=px.colors.qualitative.Plotly
        )
        st.plotly_chart(fig, use_container_width=True)
        
        if st.button("💾 Save this chart for download", key="save_chart"):
            chart_info = {
                "figure": fig,
                "chart_type": chart_type,
                "x_axis": x_axis,
                "y_axis": y_axis,
                "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
            st.session_state.saved_charts.append(chart_info)
            st.success(f"Chart saved! Total saved charts: {len(st.session_state.saved_charts)}")
        
        st.session_state.current_chart = fig
        st.session_state.chart_type = chart_type
        st.session_state.x_axis = x_axis
        st.session_state.y_axis = y_axis

except Exception as e:
    st.error(f"Error generating visualization: {str(e)}")
    st.info("Try selecting different columns or check your data types")

# ===== DATA EXPORT =====
st.divider()
st.header("Export Results")

export_format = st.selectbox(
    "Choose export format:",
    [
        "CSV (Processed Data)", 
        "CSV (Original Data)", 
        "HTML Report", 
        "Text Report", 
        "Jupyter Notebook",
        "Word Report",
        "PDF Report"
    ],
    key="export_format"
)

if export_format == "CSV (Processed Data)":
    csv = st.session_state.df.to_csv(index=False).encode('utf-8')
    st.download_button(
        "Download Processed Data (CSV)",
        csv,
        "processed_data.csv",
        help="Download filtered/processed data as CSV",
        key="csv_processed"
    )

elif export_format == "CSV (Original Data)":
    csv = st.session_state.original_df.to_csv(index=False).encode('utf-8')
    st.download_button(
        "Download Original Data (CSV)",
        csv,
        "original_data.csv",
        help="Download raw uploaded data as CSV",
        key="csv_original"
    )

elif export_format == "HTML Report":
    with st.spinner("Generating HTML report..."):
        # Create custom HTML structure with enforced colors
        html_content = f"""
        <html>
        <head>
            <title>Data Analysis Report</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #2e6c80; }}
                h2 {{ color: #3a7ca5; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .chart-container {{ margin: 30px 0; }}
            </style>
        </head>
        <body>
            <h1>Data Analysis Report</h1>
            <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            
            <h2>Data Overview</h2>
            <p><strong>Dataset:</strong> {len(st.session_state.df)} rows × {len(st.session_state.df.columns)} columns</p>
            
            <h2>Summary Statistics</h2>
            <table>
                <tr>
                    <th>Metric</th>
                    <th>Column</th>
                    <th>Value</th>
                </tr>
        """
        
        # Add properly formatted statistics table
        stats = st.session_state.df.describe().round(2)
        for col in stats.columns:
            for stat in stats.index:
                html_content += f"""
                <tr>
                    <td>{stat}</td>
                    <td>{col}</td>
                    <td>{stats.loc[stat, col]}</td>
                </tr>
                """
        html_content += "</table>"
        
        # Add saved visualizations with enforced colors
        if st.session_state.saved_charts:
            html_content += "<h2>Saved Visualizations</h2>"
            for i, chart in enumerate(st.session_state.saved_charts, 1):
                fig = chart['figure']
                # Enforce specific colors for each chart type
                if chart['chart_type'] == "Bar Chart":
                    fig.update_traces(marker_color='#1f77b4')  # Plotly blue
                elif chart['chart_type'] == "Scatter Plot":
                    fig.update_traces(marker=dict(color='#ff7f0e'))  # Plotly orange
                elif chart['chart_type'] == "Histogram":
                    fig.update_traces(marker_color='#2ca02c')  # Plotly green
                
                html_content += f"""
                <div class="chart-container">
                    <h3>Chart {i}: {chart['chart_type']} ({chart['timestamp']})</h3>
                    {fig.to_html(full_html=False, include_plotlyjs='cdn')}
                </div>
                """
        
        st.download_button(
            "Download HTML Report",
            html_content.encode('utf-8'),
            "data_analysis_report.html",
            help="Interactive HTML report with visualizations and stats",
            key="html_report"
        )

elif export_format == "Text Report":
    with st.spinner("Generating text report..."):
        text_content = f"DATA ANALYSIS REPORT\n"
        text_content += f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
        
        text_content += "DATA OVERVIEW\n"
        text_content += f"Dataset: {len(st.session_state.df)} rows × {len(st.session_state.df.columns)} columns\n\n"
        
        text_content += "SUMMARY STATISTICS\n"
        stats = st.session_state.df.describe().round(2)
        for col in stats.columns:
            text_content += f"\nColumn: {col}\n"
            for stat in stats.index:
                text_content += f"{stat}: {stats.loc[stat, col]}\n"
        
        if st.session_state.saved_charts:
            text_content += "\nSAVED VISUALIZATIONS\n"
            for i, chart in enumerate(st.session_state.saved_charts, 1):
                text_content += f"\nChart {i}: {chart['chart_type']} ({chart['timestamp']})\n"
                text_content += f"X-axis: {chart['x_axis']}\n"
                if chart['y_axis']:
                    text_content += f"Y-axis: {chart['y_axis']}\n"
        
        st.download_button(
            "Download Text Report",
            text_content.encode('utf-8'),
            "data_analysis_report.txt",
            help="Plain text report with statistics and chart details",
            key="text_report"
        )

elif export_format == "Jupyter Notebook":
    with st.spinner("Generating Jupyter notebook..."):
        # Create a new notebook
        nb = new_notebook()
        
        # Add initial cells
        nb.cells.append(new_code_cell(
            f"# Data Analysis Report\n"
            f"# Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n"
            f"print('Dataset: {len(st.session_state.df)} rows × {len(st.session_state.df.columns)} columns')"
        ))
        
        nb.cells.append(new_code_cell(
            "import pandas as pd\n"
            "import plotly.express as px\n\n"
            "# Load your data\n"
            f"df = pd.DataFrame({st.session_state.df.to_dict('list')})"
        ))
        
        # Add statistics
        nb.cells.append(new_code_cell(
            "# Summary Statistics\n"
            "df.describe()"
        ))
        
        # Add visualization cells
        if st.session_state.saved_charts:
            nb.cells.append(new_code_cell("# Saved Visualizations"))
            for i, chart in enumerate(st.session_state.saved_charts, 1):
                chart_code = f"# Chart {i}: {chart['chart_type']}\n"
                if chart['y_axis']:
                    chart_code += f"fig = px.{chart['chart_type'].lower().replace(' ', '_')}(df, x='{chart['x_axis']}', y='{chart['y_axis']}')\n"
                else:
                    chart_code += f"fig = px.{chart['chart_type'].lower().replace(' ', '_')}(df['{chart['x_axis']}'].value_counts().reset_index(), x='index', y='{chart['x_axis']}')\n"
                chart_code += "fig.show()"
                nb.cells.append(new_code_cell(chart_code))
        
        # Create download file - SIMPLIFIED AND WORKING VERSION
        notebook_content = nbformat.writes(nb)
        
        st.download_button(
            "Download Jupyter Notebook",
            notebook_content.encode('utf-8'),
            "data_analysis_report.ipynb",
            help="Jupyter notebook with analysis code",
            key="jupyter_notebook"
        )
elif export_format in ["Word Report", "PDF Report"]:
    if not st.session_state.saved_charts:
        st.warning("Please save at least one visualization before exporting reports")
        st.stop()
    
    with st.spinner("Generating professional report..."):
        try:
            os.makedirs("temp", exist_ok=True)
            
            doc = Document()
            doc.add_heading('Data Analysis Report', 0)
            
            # Data Overview
            doc.add_heading('Data Overview', level=1)
            doc.add_paragraph(f"Rows: {len(st.session_state.df)} | Columns: {len(st.session_state.df.columns)}")
            doc.add_paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Summary Statistics with proper 3-column format
            doc.add_heading('Summary Statistics', level=1)
            stats = st.session_state.df.describe().round(2)
            
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Column'
            hdr_cells[2].text = 'Value'
            
            for col in stats.columns:
                for stat in stats.index:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(stat)
                    row_cells[1].text = str(col)
                    row_cells[2].text = str(stats.loc[stat, col])
            
            # Saved Visualizations with enforced colors
            doc.add_heading('Saved Visualizations', level=1)
            for i, chart in enumerate(st.session_state.saved_charts, 1):
                chart_path = f"temp/chart_{i}.png"
                fig = chart['figure']
                
                # Apply specific colors before saving
                if chart['chart_type'] == "Bar Chart":
                    fig.update_traces(marker_color='#1f77b4')
                elif chart['chart_type'] == "Scatter Plot":
                    fig.update_traces(marker=dict(color='#ff7f0e'))
                elif chart['chart_type'] == "Histogram":
                    fig.update_traces(marker_color='#2ca02c')
                
                fig.write_image(chart_path, scale=2)
                doc.add_heading(f'Chart {i}: {chart["chart_type"]}', level=2)
                doc.add_paragraph(f"X-axis: {chart['x_axis']}")
                if chart['y_axis']:
                    doc.add_paragraph(f"Y-axis: {chart['y_axis']}")
                doc.add_picture(chart_path, width=Inches(6))
                doc.add_page_break()
            
            if export_format == "Word Report":
                # Save Word document to bytes
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                
                st.download_button(
                    "Download Word Report",
                    doc_bytes.getvalue(),
                    "data_analysis_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="word_report"
                )
            
            # PDF generation with same structure
            elif export_format == "PDF Report":
                html_content = f"""
                <html>
                <head>
                    <style>
                        body {{ font-family: Arial; }}
                        h1 {{ color: #2e6c80; }}
                        h2 {{ color: #3a7ca5; }}
                        table {{ border-collapse: collapse; width: 100%; }}
                        th, td {{ border: 1px solid #ddd; padding: 8px; }}
                        th {{ background-color: #f2f2f2; }}
                    </style>
                </head>
                <body>
                    <h1>Data Analysis Report</h1>
                    <p>Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
                    
                    <h2>Data Overview</h2>
                    <p>Rows: {len(st.session_state.df)} | Columns: {len(st.session_state.df.columns)}</p>
                    
                    <h2>Summary Statistics</h2>
                    <table>
                        <tr>
                            <th>Metric</th>
                            <th>Column</th>
                            <th>Value</th>
                        </tr>
                """
                
                for col in stats.columns:
                    for stat in stats.index:
                        html_content += f"""
                        <tr>
                            <td>{stat}</td>
                            <td>{col}</td>
                            <td>{stats.loc[stat, col]}</td>
                        </tr>
                        """
                html_content += "</table>"
                
                for i, chart in enumerate(st.session_state.saved_charts, 1):
                    chart_path = f"temp/chart_{i}.png"
                    fig = chart['figure']
                    if chart['chart_type'] == "Bar Chart":
                        fig.update_traces(marker_color='#1f77b4')
                    elif chart['chart_type'] == "Scatter Plot":
                        fig.update_traces(marker=dict(color='#ff7f0e'))
                    elif chart['chart_type'] == "Histogram":
                        fig.update_traces(marker_color='#2ca02c')
                    
                    fig.write_image(chart_path, scale=2)
                    html_content += f"""
                    <div style="page-break-before: always;">
                        <h2>Chart {i}: {chart['chart_type']}</h2>
                        <p>X-axis: {chart['x_axis']}</p>
                        {f"<p>Y-axis: {chart['y_axis']}</p>" if chart['y_axis'] else ""}
                        <img src="{chart_path}" style="width: 100%; max-width: 600px;">
                    </div>
                    """
                
                pdf_buffer = BytesIO()
                pisa.CreatePDF(html_content, dest=pdf_buffer)
                pdf_buffer.seek(0)
                
                st.download_button(
                    "Download PDF Report",
                    pdf_buffer.getvalue(),
                    "data_analysis_report.pdf",
                    mime="application/pdf",
                    key="pdf_report"
                )
            
            # Clean up temp files
            for i in range(1, len(st.session_state.saved_charts)+1):
                chart_path = f"temp/chart_{i}.png"
                if os.path.exists(chart_path):
                    os.remove(chart_path)
            
        except Exception as e:
            st.error(f"Error generating report: {str(e)}")